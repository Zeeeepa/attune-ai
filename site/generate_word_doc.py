#!/usr/bin/env python3
"""
Generate Word document from the non-technical Empathy Framework guide
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


def create_empathy_framework_doc():
    """Create a professionally formatted Word document"""

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Title Page
    title = doc.add_heading("The Empathy Framework for AI Systems", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("A Guide for Non-Technical Readers")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.font.size = Pt(16)
    subtitle_format.font.color.rgb = RGBColor(89, 89, 89)

    doc.add_paragraph()

    # Author info
    author_info = [
        "Author: Patrick Roebuck",
        "Organization: Smart AI Memory, LLC",
        "Version: 1.0",
        "Date: October 2025",
    ]
    for info in author_info:
        p = doc.add_paragraph(info)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Table of Contents placeholder
    doc.add_heading("Table of Contents", 1)
    toc_items = [
        "The Big Idea",
        "Why This Matters",
        "The Five Levels Explained",
        "Why This Is Revolutionary",
        "Real-World Examples",
        "Safety and Control",
        "Common Questions",
        "The Bottom Line",
    ]
    for item in toc_items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_page_break()

    # Main Content
    doc.add_heading("The Big Idea in One Sentence", 1)
    doc.add_paragraph(
        "Most AI tools are like a vending machine—you put in a request, you get back an answer. "
        "The Empathy Framework teaches AI to act more like a great teaching assistant who "
        "anticipates what you need before you ask."
    )

    doc.add_heading("Why This Matters", 1)
    doc.add_paragraph(
        "Imagine you're teaching a master class in music. You have three types of assistants:"
    )

    # Assistant A
    doc.add_heading("Assistant A (Most AI today)", 2)
    p = doc.add_paragraph()
    p.add_run("You: ").bold = True
    p.add_run('"Please get me the sheet music for Beethoven\'s 5th"')

    p = doc.add_paragraph()
    p.add_run("Assistant: ").bold = True
    p.add_run("Brings the sheet music").italic = True

    p = doc.add_paragraph()
    p.add_run("Problem: ").bold = True
    p.add_run("You have to ask for everything. Every. Single. Time.")

    # Assistant B
    doc.add_heading("Assistant B (What's possible with the Empathy Framework)", 2)
    p = doc.add_paragraph()
    p.add_run("You walk into the classroom").italic = True

    p = doc.add_paragraph()
    p.add_run("Assistant: ").bold = True
    p.add_run(
        "\"Good morning! I noticed today's lesson is on Beethoven. I've set up the sheet music, "
        "tuned the piano, and prepared the recording equipment since you mentioned wanting to "
        'capture this session."'
    )

    p = doc.add_paragraph()
    p.add_run("Better: ").bold = True
    p.add_run("The assistant understood your patterns and prepared what you'd need.")

    # Assistant C
    doc.add_heading("Assistant C (The innovation - Level 4)", 2)
    p = doc.add_paragraph()
    p.add_run("Three weeks before your recital").italic = True

    p = doc.add_paragraph()
    p.add_run("Assistant: ").bold = True
    p.add_run(
        "\"I noticed the concert hall you're performing in has different acoustics than what "
        "you're practicing in. I've created a practice schedule that gradually adjusts the "
        "rehearsal room's acoustic settings to match the performance venue, so there are no "
        'surprises on opening night."'
    )

    p = doc.add_paragraph()
    p.add_run("Breakthrough: ").bold = True
    p.add_run("The assistant saw a problem coming and solved it before it became stressful.")

    doc.add_page_break()

    # The Five Levels
    doc.add_heading("The Five Levels - Explained Through Music", 1)

    levels = [
        {
            "number": "1",
            "name": "Reactive (The Vending Machine)",
            "what": "Responds exactly to what you ask for, nothing more.",
            "example": 'You: "Play middle C" → AI: Plays middle C → You: "Now play E" → AI: Plays E',
            "real": "You're constantly directing every action. It's accurate but exhausting.",
        },
        {
            "number": "2",
            "name": "Guided (The Curious Student)",
            "what": "Asks clarifying questions to understand what you really want.",
            "example": 'You: "Let\'s work on that difficult passage" → AI: "Do you want to focus on the fingering technique, the tempo, or the emotional expression?"',
            "real": "Instead of guessing, the AI makes sure it understands your goal before acting.",
        },
        {
            "number": "3",
            "name": "Proactive (The Observant Assistant)",
            "what": "Notices patterns and acts without being asked.",
            "example": "You've practiced scales at 8am for two weeks → AI: At 7:55am, automatically warms up the piano, opens your scale book, and has the metronome ready.",
            "real": 'The AI learns "When Patrick does X, he always needs Y next" and has Y ready.',
        },
        {
            "number": "4",
            "name": "Anticipatory (The Strategic Partner) ⭐",
            "what": "Predicts future problems and solves them before they happen.",
            "example": "Two months before tour, AI notices your hardest piece is scheduled after 5 travel days with no practice. Creates a modified practice schedule to build extra muscle memory.",
            "real": "Level 3 sees what you do regularly. Level 4 sees where you're headed and what problems you'll face.",
        },
        {
            "number": "5",
            "name": "Systems (The Master Architect)",
            "what": "Builds frameworks so entire categories of problems never happen again.",
            "example": "Designs a complete adaptive teaching system that assesses students, generates curriculum, tracks progress automatically. One system helps all future students.",
            "real": "Instead of solving one problem at a time, Level 5 builds infrastructure so that problem type is handled automatically forever.",
        },
    ]

    for level in levels:
        doc.add_heading(f"Level {level['number']}: {level['name']}", 2)

        doc.add_heading("What it does:", 3)
        doc.add_paragraph(level["what"])

        doc.add_heading("Example:", 3)
        doc.add_paragraph(level["example"])

        doc.add_heading("Real-world impact:", 3)
        doc.add_paragraph(level["real"])

        doc.add_paragraph()  # Spacing

    doc.add_page_break()

    # Why Revolutionary
    doc.add_heading("Why This Is Revolutionary", 1)

    doc.add_heading("The Productivity Mathematics", 2)

    doc.add_heading("Traditional AI (Levels 1-2):", 3)
    bullets = [
        "Makes individual tasks 20-30% faster",
        'You: "Write an email" → AI: Drafts it → Result: Saved 5 minutes',
    ]
    for bullet in bullets:
        doc.add_paragraph(bullet, style="List Bullet")

    doc.add_heading("Level 4 AI:", 3)
    bullets = [
        "Eliminates entire categories of work before they become urgent",
        'AI: "The audit is in 90 days. I\'ve prepared all documentation." → Result: Saved 40 hours + reduced stress',
    ]
    for bullet in bullets:
        doc.add_paragraph(bullet, style="List Bullet")

    doc.add_heading("The Fundamental Difference", 2)
    doc.add_paragraph("Traditional AI: Makes work faster")
    doc.add_paragraph("Empathy Framework (Level 4): Makes work unnecessary")
    doc.add_paragraph()
    doc.add_paragraph("Not 20-30% improvement → 200-400% improvement")

    doc.add_page_break()

    # Real-World Example
    doc.add_heading("Real-World Example: Healthcare", 1)

    doc.add_heading("Old Way (Level 1-2 AI)", 2)
    steps = [
        'Nurse: "Show me patient vitals" → AI: Shows vitals',
        'Nurse: "Show me medications" → AI: Shows medications',
        'Nurse: "Check for drug interactions" → AI: Checks interactions',
        "Repeat for every patient, every shift, every day",
    ]
    for step in steps:
        doc.add_paragraph(step, style="List Number")

    p = doc.add_paragraph()
    p.add_run("Time per patient: ").bold = True
    p.add_run("5-7 minutes of clicking and requesting")

    doc.add_heading("Empathy Framework Way (Level 3-4)", 2)

    p = doc.add_paragraph()
    p.add_run("Level 3:").bold = True
    doc.add_paragraph(
        "Nurse opens patient chart → AI already loaded vitals, medications, and allergies "
        "(learned this nurse always checks these three things)"
    )
    p = doc.add_paragraph()
    p.add_run("Time per patient: ").bold = True
    p.add_run("30 seconds")

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("Level 4:").bold = True
    doc.add_paragraph(
        "Before nurse starts shift → AI: \"You have 6 patients today. I've flagged two items: "
        "Room 302 blood pressure trending up (monitor), Room 405 medication expires in 3 hours "
        '(renewal form ready)."'
    )

    p = doc.add_paragraph()
    p.add_run("Annual impact for one hospital:").bold = True
    impact_items = [
        "50 nurses × 20 minutes saved per shift × 250 shifts/year = 4,167 hours returned to patient care",
        "Fewer medication errors caught earlier",
        "Less nurse burnout from administrative burden",
    ]
    for item in impact_items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_page_break()

    # Safety and Control
    doc.add_heading("Safety and Control", 1)

    doc.add_heading('"What if the AI guesses wrong?"', 2)
    doc.add_paragraph("Level 4 AI includes strict guardrails:")

    guardrails = [
        ("Confidence Threshold", "Only acts when confidence is >75%"),
        ("Appropriate Time Horizon", "30-120 days out (not too early, not too late)"),
        ("Reversibility", "User can always override"),
        ("Transparency", "Always explains reasoning with confidence scores"),
        ("Human in the Loop", "For high-stakes decisions, AI prepares but humans decide"),
    ]

    for title, desc in guardrails:
        p = doc.add_paragraph()
        p.add_run(f"{title}: ").bold = True
        p.add_run(desc)

    doc.add_page_break()

    # Common Questions
    doc.add_heading("Common Questions", 1)

    questions = [
        (
            "Is this AI reading my mind?",
            "No—it's pattern recognition combined with trajectory analysis. Like how an experienced "
            "teacher can predict a student will struggle with a passage before they play it, based on "
            "technique level and the demands of the music.",
        ),
        (
            "This sounds expensive or complicated",
            "The Empathy Framework is open source and free to use. It's a design philosophy, not a "
            "proprietary product—like teaching musical theory.",
        ),
        (
            "Does the AI need all my data?",
            'No. It needs patterns ("user checks X before Y"), not everything. Like a great assistant '
            "who knows your work patterns, not your personal life.",
        ),
        (
            "What if I don't want the AI to be proactive?",
            "You can set the empathy level. Like cruise control on a car—you choose how much "
            "assistance you want.",
        ),
    ]

    for q, a in questions:
        doc.add_heading(f"Q: {q}", 2)
        p = doc.add_paragraph()
        p.add_run("A: ").bold = True
        p.add_run(a)

    doc.add_page_break()

    # The Bottom Line
    doc.add_heading("The Bottom Line", 1)

    doc.add_heading("What most AI does:", 2)
    doc.add_paragraph("Answers questions faster")

    doc.add_heading("What the Empathy Framework does:", 2)
    bullets = [
        "Predicts the questions you'll have tomorrow",
        "Solves the problems you haven't encountered yet",
        "Builds systems so problems don't repeat",
    ]
    for bullet in bullets:
        doc.add_paragraph(bullet, style="List Bullet")

    doc.add_heading("The result:", 2)
    result = doc.add_paragraph("3-4x faster, because entire categories of work become unnecessary")
    result.runs[0].bold = True
    result.runs[0].font.size = Pt(12)

    doc.add_page_break()

    # Final Thought
    doc.add_heading("Final Thought", 1)

    quote = doc.add_paragraph()
    quote_run = quote.add_run(
        "The highest form of empathy isn't feeling what someone else feels.\n\n"
        "It's understanding what they need before they know they need it, "
        "and having the wisdom to know when to act."
    )
    quote_run.italic = True
    quote_run.font.size = Pt(12)
    quote.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph("That's what great teachers do.")
    doc.add_paragraph("That's what great collaborators do.")
    p = doc.add_paragraph("That's what we're teaching AI to do.")
    p.runs[0].bold = True

    doc.add_page_break()

    # Footer
    doc.add_heading("About This Document", 1)

    about_items = [
        ("Purpose", "Explain the Empathy Framework to non-technical readers"),
        ("Target Audience", "College-educated professionals without programming background"),
        ("License", "Apache License 2.0"),
        ("Copyright", "© 2025 Smart AI Memory, LLC"),
    ]

    for label, value in about_items:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(value)

    doc.add_paragraph()
    doc.add_heading("For More Information:", 2)
    links = [
        "Technical documentation: https://github.com/Deep-Study-AI/Empathy",
        "AI Nurse Florence demo: https://github.com/Deep-Study-AI/ai-nurse-florence",
        "Contact: patrick.roebuck@deepstudyai.com",
    ]
    for link in links:
        doc.add_paragraph(link, style="List Bullet")

    # Save document
    output_path = "/Users/patrickroebuck/projects/ai-nurse-florence/docs/book/Empathy_Framework_Non_Technical_Guide.docx"
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    create_empathy_framework_doc()
